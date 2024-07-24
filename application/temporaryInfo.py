import datetime
import os

from Printer import print_file
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
class temp_info:
    def __init__(self, DN=None, thaiTitle=None, thaiName=None, thaiSurname=None, 
                    engTitle=None, engName=None, engSurname=None, 
                    ID=None, birthdate = None, job = None, age=None, gender=None, status=None, 
                    address=None, phoneNum=None,
                    emergencyName=None, relation=None,
                    emergencyNum=None, emergencyAddress=None):
        self.DN = DN
        self.thaiTitle = thaiTitle
        self.thaiName = thaiName
        self.thaiSurname = thaiSurname
        self.engTitle = engTitle
        self.engName = engName
        self.engSurname = engSurname
        self.ID = ID
        self.birthdate = birthdate
        self.job = job
        self.age = age
        self.gender = gender
        self.status = status
        self.address = address
        self.phoneNum = phoneNum
        self.emergencyName = emergencyName
        self.relation = relation
        self.emergencyNum = emergencyNum
        self.emergencyAddress = emergencyAddress
    # ฟังก์ชันทำใบคัดกรอง และทำ.docx
    def makeDocument(self, thaiTitle, thaiName, thaiSurname, ID,
                    tempVar, bloodtopVar, bloodbottomVar, hearrateVar, 
                    checkbox1Var, checkbox2Var, checkbox3Var, checkbox4Var,
                    checkbox5Var, checkbox6Var, checkbox7Var, checkbox8Var,
                    checkbox9Var, checkbox10Var, checkbox11Var, checkbox12Var,
                    checkbox13Var, checkbox14Var, checkbox15Var, checkbox15inputVar,
                    checkbox16Var, checkbox17Var, checkbox18Var, checkbox19Var,
                    checkbox20Var, checkbox21Var, checkbox22Var, checkbox23Var, 
                    checkbox24Var, checkbox25Var, checkbox26Var, checkbox27Var,
                    checkbox28Var, checkbox29Var, checkbox30Var, checkbox30inputVar):
        self.thaiTitle = thaiTitle
        self.thaiName = thaiName
        self.thaiSurname = thaiSurname
        self.ID = ID
        self.temp_var = tempVar
        self.bloodtop_var = bloodtopVar
        self.bloodbottom_var = bloodbottomVar
        self.hearrate_var = hearrateVar
        self.checkbox1_var = checkbox1Var
        self.checkbox2_var = checkbox2Var
        self.checkbox3_var = checkbox3Var
        self.checkbox4_var = checkbox4Var
        self.checkbox5_var = checkbox5Var
        self.checkbox6_var = checkbox6Var
        self.checkbox7_var = checkbox7Var
        self.checkbox8_var = checkbox8Var
        self.checkbox9_var = checkbox9Var
        self.checkbox10_var = checkbox10Var
        self.checkbox11_var = checkbox11Var
        self.checkbox12_var = checkbox12Var
        self.checkbox13_var = checkbox13Var
        self.checkbox14_var = checkbox14Var
        self.checkbox15_var = checkbox15Var
        self.checkbox15input_var = checkbox15inputVar
        self.checkbox16_var = checkbox16Var
        self.checkbox17_var = checkbox17Var
        self.checkbox18_var = checkbox18Var
        self.checkbox19_var = checkbox19Var
        self.checkbox20_var = checkbox20Var
        self.checkbox21_var = checkbox21Var
        self.checkbox22_var = checkbox22Var
        self.checkbox23_var = checkbox23Var
        self.checkbox24_var = checkbox24Var
        self.checkbox25_var = checkbox25Var
        self.checkbox26_var = checkbox26Var
        self.checkbox27_var = checkbox27Var
        self.checkbox28_var = checkbox28Var
        self.checkbox29_var = checkbox29Var
        self.checkbox30_var = checkbox30Var
        self.checkbox30input_var = checkbox30inputVar
        # เดือนไทย
        def month_thai(month_num):
            months = [
                "ม.ค.",
                "ก.พ.",
                "มี.ค.",
                "เม.ย.",
                "พ.ค.",
                "มิ.ย.",
                "ก.ค",
                "ส.ค.",
                "ก.ย.",
                "ต.ค",
                "พ.ย.",
                "ธ.ค.",
            ]
            return months[month_num - 1]
        ## สร้างไฟล์ .docx สำหรับทำใบคัดกรอง
        # สร้างตัวแปรไว้เก็บแสดงวันที่
        now = datetime.datetime.now()
        buddhistYear = now.year + 543
        thai_month = month_thai(now.month)
        date = f"{now.day} {thai_month} {buddhistYear}"
        
        # อ่านจากไฟล์ต้นฉบับ
        doc = Document("components/Paperform.docx")
        # รูปติ๊กถูก
        checkImg = "components/check.png"
    
        ### กล่องใส่วันที่ใน .doc / table 0
        dateDoc = doc.tables[0].cell(0,1)
        datePara = dateDoc.paragraphs[0]
        datePara.clear()
        
        newdateDoc = datePara.add_run(f"{date}")
        fontforDate = newdateDoc.font
        fontforDate.name = "FreesiaUPC"
        fontforDate.size = Pt(14)
        newdateDoc.element.rPr.rFonts.set(qn('w:eastAsia'), fontforDate.name)
        
        ### กล่องใส่ ชื่อ อุณหภูมิ ผลความดันโลหิต / table 1
        ## ชื่อ
        nameDoc = doc.tables[1].cell(1,1)
        namePara = nameDoc.paragraphs[0]
        namePara.clear()
        
        newnameDoc = namePara.add_run(f"{self.thaiTitle}{self.thaiName} {self.thaiSurname}")
        fontforName = newnameDoc.font
        fontforName.name = "FreesiaUPC"
        fontforName.size = Pt(14)
        newnameDoc.element.rPr.rFonts.set(qn('w:eastAsia'), fontforName.name)
        ## ผลความดันโลหิต
        sysDoc = doc.tables[1].cell(2,3)
        sysPara = sysDoc.paragraphs[0]
        sysPara.clear()
        
        newsysDoc = sysPara.add_run(f"{self.bloodtop_var}")
        fontforSys = newsysDoc.font
        fontforSys.name = "FreesiaUPC"
        fontforSys.size = Pt(14)
        newsysDoc.element.rPr.rFonts.set(qn('w:eastAsia'), fontforSys.name)
        
        diaDoc = doc.tables[1].cell(2,6)
        diaPara = diaDoc.paragraphs[0]
        diaPara.clear()
        
        newdiaDoc = diaPara.add_run(f"{self.bloodbottom_var}")
        fontforDia = newdiaDoc.font
        fontforDia.name = "FreesiaUPC"
        fontforDia.size = Pt(14)
        newdiaDoc.element.rPr.rFonts.set(qn('w:eastAsia'), fontforDia.name)
        
        prDoc = doc.tables[1].cell(3,3)
        prPara = prDoc.paragraphs[0]
        prPara.clear()
        
        newprDoc = prPara.add_run(f"{self.hearrate_var}")
        fontforPr = newprDoc.font
        fontforPr.name = "FreesiaUPC"
        fontforPr.size = Pt(14)
        newprDoc.element.rPr.rFonts.set(qn('w:eastAsia'), fontforPr.name)
        
        tempDoc = doc.tables[1].cell(3,6)
        tempPara = tempDoc.paragraphs[0]
        tempPara.clear()
        
        newtempDoc = tempPara.add_run(f"{self.temp_var}")
        fontforTemp = newtempDoc.font
        fontforTemp.name = "FreesiaUPC"
        fontforTemp.size = Pt(14)
        newtempDoc.element.rPr.rFonts.set(qn('w:eastAsia'), fontforTemp.name)
        
        ## กล่องใส่ข้อมูลคัดกรองอาการ / table 2
        ## โรคไข้หวัด
        # ไม่มีอาการเหล่านี้
        if checkbox1Var == 1:
            checkbox1Doc = doc.tables[2].cell(2,0)
            checkbox1Para = checkbox1Doc.paragraphs[0]
            checkbox1Para.clear()
            
            runforcheckbox1 = checkbox1Para.add_run()
            runforcheckbox1.add_picture(checkImg, width=Inches(0.32))
            newcheckbox1 = checkbox1Para.add_run("ไม่มีอาการ")
            fontforcheckbox1 = newcheckbox1.font
            fontforcheckbox1.name = "FreesiaUPC"
            fontforcheckbox1.size = Pt(14)
            newcheckbox1.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox1.name)
        # มีไข้ > 37.5
        if checkbox2Var == 1:
            checkbox2Para = checkbox1Doc.paragraphs[1]
            checkbox2Para.clear()
            
            runforcheckbox2 = checkbox2Para.add_run()
            runforcheckbox2.add_picture(checkImg, width=Inches(0.32))
            newcheckbox2 = checkbox2Para.add_run("1.มีไข้ (อุณหภูมิ > 37.5 °C)")
            fontforcheckbox2 = newcheckbox2.font
            fontforcheckbox2.name = "FreesiaUPC"
            fontforcheckbox2.size = Pt(14)
            newcheckbox2.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox2.name)
        # ไอ จาม มีน้ำมูก
        if checkbox3Var == 1:
            checkbox3Para = checkbox1Doc.paragraphs[2]
            checkbox3Para.clear()
            
            runforcheckbox3 = checkbox3Para.add_run()
            runforcheckbox3.add_picture(checkImg, width=Inches(0.32))
            newcheckbox3 = checkbox3Para.add_run("2.ไอ จาม มีน้ำมูก")
            fontforcheckbox3 = newcheckbox3.font
            fontforcheckbox3.name = "FreesiaUPC"
            fontforcheckbox3.size = Pt(14)
            newcheckbox3.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox3.name)
        # มีเสมหะ เจ็บคอ
        if checkbox4Var == 1:
            checkbox4Para = checkbox1Doc.paragraphs[3]
            checkbox4Para.clear()
            
            runforcheckbox4 = checkbox4Para.add_run()
            runforcheckbox4.add_picture(checkImg, width=Inches(0.32))
            newcheckbox4 = checkbox4Para.add_run("3.มีเสมหะ เจ็บคอ")
            fontforcheckbox4 = newcheckbox4.font
            fontforcheckbox4.name = "FreesiaUPC"
            fontforcheckbox4.size = Pt(14)
            newcheckbox4.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox4.name)
        # ปวดศีรษะ
        if checkbox5Var == 1:
            checkbox5Para = checkbox1Doc.paragraphs[4]
            checkbox5Para.clear()
            
            runforcheckbox5 = checkbox5Para.add_run()
            runforcheckbox5.add_picture(checkImg, width=Inches(0.32))
            newcheckbox5 = checkbox5Para.add_run("ปวดศีรษะ")
            fontforcheckbox5 = newcheckbox5.font
            fontforcheckbox5.name = "FreesiaUPC"
            fontforcheckbox5.size = Pt(14)
            newcheckbox5.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox5.name)
        # มีอ่อนเพลีย
        if checkbox6Var == 1:
            checkbox6Para = checkbox1Doc.paragraphs[5]
            checkbox6Para.clear()
            
            runforcheckbox6 = checkbox6Para.add_run()
            runforcheckbox6.add_picture(checkImg, width=Inches(0.32))
            newcheckbox6 = checkbox6Para.add_run("5.มีอ่อนเพลีย")
            fontforcheckbox6 = newcheckbox6.font
            fontforcheckbox6.name = "FreesiaUPC"
            fontforcheckbox6.size = Pt(14)
            newcheckbox6.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox6.name)
        # ปวดกล้ามเนื้อ
        if checkbox7Var == 1:
            checkbox7Para = checkbox1Doc.paragraphs[6]
            checkbox7Para.clear()
            
            runforcheckbox7 = checkbox7Para.add_run()
            runforcheckbox7.add_picture(checkImg, width=Inches(0.32))
            newcheckbox7 = checkbox7Para.add_run("6.ปวดกล้ามเนื้อ")
            fontforcheckbox7 = newcheckbox7.font
            fontforcheckbox7.name = "FreesiaUPC"
            fontforcheckbox7.size = Pt(14)
            newcheckbox7.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox7.name)
        ## โรควัณโรค
        # ไม่มีอาการ
        if checkbox8Var == 1:
            checkbox8Doc = doc.tables[2].cell(2,1)
            checkbox8Para = checkbox8Doc.paragraphs[0]
            checkbox8Para.clear()
            
            runforcheckbox8 = checkbox8Para.add_run()
            runforcheckbox8.add_picture(checkImg, width=Inches(0.32))
            newcheckbox8 = checkbox8Para.add_run("ไม่มีอาการ")
            fontforcheckbox8 = newcheckbox8.font
            fontforcheckbox8.name = "FreesiaUPC"
            fontforcheckbox8.size = Pt(14)
            newcheckbox8.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox8.name)
        # ไอเรื้อรังเกิน 2 สัปดาห์
        if checkbox9Var == 1:
            checkbox9Para = checkbox8Doc.paragraphs[1]
            checkbox9Para.clear()
            
            runforcheckbox9 = checkbox9Para.add_run()
            runforcheckbox9.add_picture(checkImg, width=Inches(0.32))
            newcheckbox9 = checkbox9Para.add_run("1.ไอเรื้อรังเกิน 2 สัปดาห์")
            fontforcheckbox9 = newcheckbox9.font
            fontforcheckbox9.name = "FreesiaUPC"
            fontforcheckbox9.size = Pt(14)
            newcheckbox9.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox9.name)
        # ไอมีเลือดปน
        if checkbox10Var == 1:
            checkbox10Para = checkbox8Doc.paragraphs[2]
            checkbox10Para.clear()
            
            runforcheckbox10 = checkbox10Para.add_run()
            runforcheckbox10.add_picture(checkImg, width=Inches(0.32))
            newcheckbox10 = checkbox10Para.add_run("2.ไอมีเลือดปน")
            fontforcheckbox10 = newcheckbox10.font
            fontforcheckbox10.name = "FreesiaUPC"
            fontforcheckbox10.size = Pt(14)
            newcheckbox10.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox10.name)
        # น้ำหนักลด 3-5 กก/เดือน
        if checkbox11Var == 1:
            checkbox11Para = checkbox8Doc.paragraphs[3]
            checkbox11Para.clear()
            
            runforcheckbox11 = checkbox11Para.add_run()
            runforcheckbox11.add_picture(checkImg, width=Inches(0.32))
            newcheckbox11 = checkbox11Para.add_run("3.น้ำหนักลด 3-5 กก./เดือนโดย")
            fontforcheckbox11 = newcheckbox11.font
            fontforcheckbox11.name = "FreesiaUPC"
            fontforcheckbox11.size = Pt(14)
            newcheckbox11.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox11.name)
        # ไข้ตอนบ่ายเกิน 2 สัปดาห์
        if checkbox12Var == 1:
            checkbox12Para = checkbox8Doc.paragraphs[5]
            checkbox12Para.clear()

            runforcheckbox12 = checkbox12Para.add_run()
            runforcheckbox12.add_picture(checkImg, width=Inches(0.32))
            newcheckbox12 = checkbox12Para.add_run("4.ไข้ตอนบ่ายเกิน 2 สัปดาห์")
            fontforcheckbox12 = newcheckbox12.font
            fontforcheckbox12.name = "FreesiaUPC"
            fontforcheckbox12.size = Pt(14)
            newcheckbox12.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox12.name)
        # มีเหงื่อออกกลางคืนใน 1 เดือน
        if checkbox13Var == 1:
            checkbox13Para = checkbox8Doc.paragraphs[6]
            checkbox13Para.clear()

            runforcheckbox13 = checkbox13Para.add_run()
            runforcheckbox13.add_picture(checkImg, width=Inches(0.32))
            newcheckbox13 = checkbox13Para.add_run("5.มีเหงื่อออกกลางคืนใน 1 เดือน")
            fontforcheckbox13 = newcheckbox13.font
            fontforcheckbox13.name = "FreesiaUPC"
            fontforcheckbox13.size = Pt(14)
            newcheckbox13.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox13.name)
        # มีประวัติสัมผัสกับผู้ป่วยวัณโรค
        if checkbox14Var == 1:
            checkbox14Para = checkbox8Doc.paragraphs[7]
            checkbox14Para.clear()

            runforcheckbox14 = checkbox14Para.add_run()
            runforcheckbox14.add_picture(checkImg, width=Inches(0.32))
            newcheckbox14 = checkbox14Para.add_run("6.มีประวัติสัมผัสกับผู้ป่วยวัณโรค")
            fontforcheckbox14 = newcheckbox14.font
            fontforcheckbox14.name = "FreesiaUPC"
            fontforcheckbox14.size = Pt(14)
            newcheckbox14.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox14.name)
        # กำลังรักษาโรควัณโรค
        # print(doc.tables[2].cell(2,1).paragraphs[9].text) เอาไว้ข้อมูลจาก word
        if checkbox15Var == 1:
            checkbox15Para = checkbox8Doc.paragraphs[8]
            checkbox15Para.clear()

            runforcheckbox15 = checkbox15Para.add_run()
            runforcheckbox15.add_picture(checkImg, width=Inches(0.32))
            newcheckbox15 = checkbox15Para.add_run("7.กำลังรักษาโรควัณโรค")
            fontforcheckbox15 = newcheckbox15.font
            fontforcheckbox15.name = "FreesiaUPC"
            fontforcheckbox15.size = Pt(14)
            newcheckbox15.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox15.name)
            
            checkbox15inputPara = checkbox8Doc.paragraphs[9]
            checkbox15inputPara.clear()
            
            newcheckbox15input = checkbox15inputPara.add_run(f"(ระยะเวลาในการรักษา {self.checkbox15input_var}   )")
            fontforcheckbox15input = newcheckbox15input.font
            fontforcheckbox15input.name = "FreesiaUPC"
            fontforcheckbox15input.size = Pt(14)
            newcheckbox15.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox15input.name)
        # เคยมีประวัติเป็นโรควัณโรคและมีใบรับรองแพทย์
        if checkbox16Var == 1:
            checkbox16Doc = doc.tables[2].cell(3,1)
            checkbox16Para = checkbox16Doc.paragraphs[2]
            checkbox16Para.clear()

            runforcheckbox16 = checkbox16Para.add_run()
            runforcheckbox16.add_picture(checkImg, width=Inches(0.32))
            newcheckbox16 = checkbox16Para.add_run("เคยมีประวัติเป็นโรควัณโรคและมีใบรับรองแพทย์")
            fontforcheckbox16 = newcheckbox16.font
            fontforcheckbox16.name = "FreesiaUPC"
            fontforcheckbox16.size = Pt(14)
            newcheckbox16.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox16.name)
        # โรคเริมและงูสวัด
        # ไม่มีอาการ
        if checkbox17Var == 1:
            checkbox17Doc = doc.tables[2].cell(3,2)
            checkbox17Para = checkbox17Doc.paragraphs[0]
            checkbox17Para.clear()

            runforcheckbox17 = checkbox17Para.add_run()
            runforcheckbox17.add_picture(checkImg, width=Inches(0.32))
            newcheckbox17 = checkbox17Para.add_run("ไม่มีอาการ")
            fontforcheckbox17 = newcheckbox17.font
            fontforcheckbox17.name = "FreesiaUPC"
            fontforcheckbox17.size = Pt(14)
            newcheckbox17.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox17.name)
        # มีตุ่มน้ำที่ริมฝีปาก
        if checkbox18Var == 1:
            checkbox18Para = checkbox17Doc.paragraphs[1]
            checkbox18Para.clear()

            runforcheckbox18 = checkbox18Para.add_run()
            runforcheckbox18.add_picture(checkImg, width=Inches(0.32))
            newcheckbox18 = checkbox18Para.add_run("1.มีตุ่มน้ำที่ริมฝีปาก")
            fontforcheckbox18 = newcheckbox18.font
            fontforcheckbox18.name = "FreesiaUPC"
            fontforcheckbox18.size = Pt(14)
            newcheckbox18.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox18.name)
        # แผลที่มีอาการเจ็บแสบร้อนที่ริมฝีปาก
        if checkbox19Var == 1:
            checkbox19Para = checkbox17Doc.paragraphs[2]
            checkbox19Para.clear()

            runforcheckbox19 = checkbox19Para.add_run()
            runforcheckbox19.add_picture(checkImg, width=Inches(0.32))
            newcheckbox19 = checkbox19Para.add_run("2.แผลที่มีอาการเจ็บแสบร้อน")
            fontforcheckbox19 = newcheckbox19.font
            fontforcheckbox19.name = "FreesiaUPC"
            fontforcheckbox19.size = Pt(14)
            newcheckbox19.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox19.name)
        # มีตุ่มน้ำใสเป็นแนวยาวจามผิวหนังร่างกาย
        if checkbox20Var == 1:
            checkbox20Para = checkbox17Doc.paragraphs[4]
            checkbox20Para.clear()

            runforcheckbox20 = checkbox20Para.add_run()
            runforcheckbox20.add_picture(checkImg, width=Inches(0.32))
            newcheckbox20 = checkbox20Para.add_run("3.มีตุ่มน้ำใสเป็นแนวยาวตาม")
            fontforcheckbox20 = newcheckbox20.font
            fontforcheckbox20.name = "FreesiaUPC"
            fontforcheckbox20.size = Pt(14)
            newcheckbox20.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox20.name)
        # รู้สึกเจ็บแปลบบริเวณผิวหนัง
        if checkbox21Var == 1:
            checkbox21Para = checkbox17Doc.paragraphs[6]
            checkbox21Para.clear()

            runforcheckbox21 = checkbox21Para.add_run()
            runforcheckbox21.add_picture(checkImg, width=Inches(0.32))
            newcheckbox21 = checkbox21Para.add_run("4.รู้สึกเจ็บแปลบบริเวณผิวหนัง")
            fontforcheckbox21 = newcheckbox21.font
            fontforcheckbox21.name = "FreesiaUPC"
            fontforcheckbox21.size = Pt(14)
            newcheckbox21.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox21.name)
        # รู้สึกคัน ปวดแสบ ปวดร้อนบริเวณผิวหนัง
        if checkbox22Var == 1:
            checkbox22Para = checkbox17Doc.paragraphs[7]
            checkbox22Para.clear()

            runforcheckbox22 = checkbox22Para.add_run()
            runforcheckbox22.add_picture(checkImg, width=Inches(0.32))
            newcheckbox22 = checkbox22Para.add_run("5.รู้สึกคัน ปวดแสบ ปวดร้อน")
            fontforcheckbox22 = newcheckbox22.font
            fontforcheckbox22.name = "FreesiaUPC"
            fontforcheckbox22.size = Pt(14)
            newcheckbox22.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox22.name)
        # มีประวัติเคยเป็นเริมหรืองูสวัด
        if checkbox23Var == 1:
            checkbox23Para = checkbox17Doc.paragraphs[9]
            checkbox23Para.clear()

            runforcheckbox23 = checkbox23Para.add_run()
            runforcheckbox23.add_picture(checkImg, width=Inches(0.32))
            newcheckbox23 = checkbox23Para.add_run("6.มีประวัติเคยเป็นเริมหรืองูสวัด")
            fontforcheckbox23 = newcheckbox23.font
            fontforcheckbox23.name = "FreesiaUPC"
            fontforcheckbox23.size = Pt(14)
            newcheckbox23.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox23.name)
        ### แบบคัดกรองโรคไม่ติดเชื้อ / table 3
        ## โรคความดันโลหิตสูง
        if checkbox24Var == 1:
            checkbox24Doc = doc.tables[3].cell(1,0)
            checkbox24Para = checkbox24Doc.paragraphs[0]
            checkbox24Para.clear()
            
            runforcheckbox24 = checkbox24Para.add_run()
            runforcheckbox24.add_picture(checkImg, width=Inches(0.32))
            newcheckbox24 = checkbox24Para.add_run("โรคความดันโลหิตสูง")
            fontforcheckbox24 = newcheckbox24.font
            fontforcheckbox24.name = "FreesiaUPC"
            fontforcheckbox24.size = Pt(14)
            newcheckbox24.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox24.name)
        ## โรคเบาหวาน
        if checkbox25Var == 1:
            checkbox25Doc = doc.tables[3].cell(1,3)
            checkbox25Para = checkbox25Doc.paragraphs[0]
            checkbox25Para.clear()
            
            runforcheckbox25 = checkbox25Para.add_run()
            runforcheckbox25.add_picture(checkImg, width=Inches(0.32))
            newcheckbox25 = checkbox25Para.add_run("โรคเบาหวาน")
            fontforcheckbox25 = newcheckbox25.font
            fontforcheckbox25.name = "FreesiaUPC"
            fontforcheckbox25.size = Pt(14)
            newcheckbox25.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox25.name)
        ## โรคหัวใจ
        if checkbox26Var == 1:
            checkbox26Doc = doc.tables[3].cell(1,4)
            checkbox26Para = checkbox26Doc.paragraphs[0]
            checkbox26Para.clear()
            
            runforcheckbox26 = checkbox26Para.add_run()
            runforcheckbox26.add_picture(checkImg, width=Inches(0.32))
            newcheckbox26 = checkbox26Para.add_run("โรคหัวใจ")
            fontforcheckbox26 = newcheckbox26.font
            fontforcheckbox26.name = "FreesiaUPC"
            fontforcheckbox26.size = Pt(14)
            newcheckbox26.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox26.name)
        ## โรคไทรอยด์
        if checkbox27Var == 1:
            checkbox27Doc = doc.tables[3].cell(1,5)
            checkbox27Para = checkbox27Doc.paragraphs[0]
            checkbox27Para.clear()
            
            runforcheckbox27 = checkbox27Para.add_run()
            runforcheckbox27.add_picture(checkImg, width=Inches(0.32))
            newcheckbox27 = checkbox27Para.add_run("โรคไทรอยด์")
            fontforcheckbox27 = newcheckbox27.font
            fontforcheckbox27.name = "FreesiaUPC"
            fontforcheckbox27.size = Pt(14)
            newcheckbox27.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox27.name)
        ## เคยมีประวัติเป็นโรคหลอดเลือดสมอง (stroke) หรือเคยมีอาการ
        if checkbox28Var == 1:
            checkbox28Doc = doc.tables[3].cell(2,0)
            checkbox28Para = checkbox28Doc.paragraphs[0]
            checkbox28Para.clear()
            
            runforcheckbox28 = checkbox28Para.add_run()
            runforcheckbox28.add_picture(checkImg, width=Inches(0.32))
            newcheckbox28 = checkbox28Para.add_run("เคยมีประวัติเป็นโรคหลอดเลือดสมอง (stroke) หรือเคยมีอาการ")
            fontforcheckbox28 = newcheckbox28.font
            fontforcheckbox28.name = "FreesiaUPC"
            fontforcheckbox28.size = Pt(14)
            newcheckbox28.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox28.name)
        ## โรคภูมิคุ้มกันบกพร่อง
        if checkbox29Var == 1:
            checkbox29Doc = doc.tables[3].cell(2,4)
            checkbox29Para = checkbox29Doc.paragraphs[0]
            checkbox29Para.clear()
            
            runforcheckbox29 = checkbox29Para.add_run()
            runforcheckbox29.add_picture(checkImg, width=Inches(0.32))
            newcheckbox29 = checkbox29Para.add_run("โรคภูมิคุ้มกันบกพร่อง")
            fontforcheckbox29 = newcheckbox29.font
            fontforcheckbox29.name = "FreesiaUPC"
            fontforcheckbox29.size = Pt(14)
            newcheckbox29.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox29.name)
        ## สตรีมีครรภ์
        if checkbox30Var == 1:
            checkbox30Doc = doc.tables[3].cell(3,0)
            checkbox30Para = checkbox30Doc.paragraphs[0]
            checkbox30Para.clear()
            
            runforcheckbox30 = checkbox30Para.add_run()
            runforcheckbox30.add_picture(checkImg, width=Inches(0.32))
            newcheckbox30 = checkbox30Para.add_run("สตรีมีครรภ์ อายุครรภ์")
            fontforcheckbox30 = newcheckbox30.font
            fontforcheckbox30.name = "FreesiaUPC"
            fontforcheckbox30.size = Pt(14)
            newcheckbox30.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox30.name)
            
            checkbox30inputDoc = doc.tables[3].cell(3,1)
            checkbox30inputPara = checkbox30inputDoc.paragraphs[0]
            checkbox30inputPara.clear()
            
            newcheckbox30input = checkbox30inputPara.add_run(f"{self.checkbox30input_var}")
            fontforcheckbox30input = newcheckbox30input.font
            fontforcheckbox30input.name = "FreesiaUPC"
            fontforcheckbox30input.size = Pt(14)
            newcheckbox30input.element.rPr.rFonts.set(qn('w:eastAsia'), fontforcheckbox30input.name)
        
        # สร้าง folder 
        os.makedirs("documents", exist_ok=True)
        
        # เซฟไฟล์ .docx
        docfile = f"{self.ID}_{self.thaiName}.docx"
        doc.save(f"documents/{docfile}")
    
        # สั่งพิมพ์ใบคัดกรอง
        # print_file(f"documents/{docfile}")
        
# tempInfo = temp_info()
# tempInfo.makeDocument(None, None, None, None, None, None, None, None, None, None, None, None
#                       , None, None, None, None, None, None, None, None, None, None, None, None
#                       , None, None, None, None, None, None, None, None, None, None, None, None
#                       , None, None, None, None, None)